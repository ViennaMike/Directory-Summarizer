# -*- coding: utf-8 -*-
"""
Created on Thu Jun 04-06 21:24:16 2015

@author: Mike McGurrin
"""

# for sumy
from __future__ import absolute_import
from __future__ import division, print_function, unicode_literals

#from sumy.parsers.html import HtmlParser # Keep in case add html later
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer as Summarizer
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words

import os, os.path
import errno
import sys
reload(sys)
sys.setdefaultencoding("utf-8")

import pdf2txt
import docx2txt
from docx import Document
#import docx2txt
import textract

def main(sum_type, startpath, fileList, destination, length): 
    
    # Safe opening whether or not path exists
    # Taken from http://stackoverflow.com/a/600612/119527
    def mkdir_p(path):
        """ Create directory if needed"""
        try:
            os.makedirs(path)
        except OSError as exc: # Python >2.5
            if exc.errno == errno.EEXIST and os.path.isdir(path):
                pass
            else: raise
    
    def safe_open_a(path):
        ''' Open "path" for writing, creating any parent directories as needed.
        '''
        mkdir_p(os.path.dirname(path))
        return open(path, 'a')
    
    def recursive_glob(rootdir='.', suffix=()):
        """ recursively traverses full path from route, returns
            paths and file names for files with suffix in tuple """
        pathlist = []
        filelist = []
        for looproot,dirnames, filenames in os.walk(rootdir):
            for filename in filenames:
                if filename.endswith(suffix):
                    pathlist.append(os.path.join(looproot, filename))
                    filelist.append(filename)
        return pathlist, filelist
        
    def valid_xml_char_ordinal(c):
        codepoint = ord(c)
        # conditions ordered by presumed frequency
        return (
            0x20 <= codepoint <= 0xD7FF or
            codepoint in (0x9, 0xA, 0xD) or
            0xE000 <= codepoint <= 0xFFFD or
            0x10000 <= codepoint <= 0x10FFFF
            )

    def use_sumy(input, SENTENCES_COUNT, method, parser_option):
        """Code to run sumy
        # Supported summarization methods:
        #    Luhn - heurestic method, reference
        #    Edmundson heurestic method with previous statistic research, reference
        #    Latent Semantic Analysis, LSA - one of the algorithm from http://scholar.google.com/citations?user=0fTuW_YAAAAJ&hl=en I think the author is using more advanced algorithms now. Steinberger, J. a Ježek, K. Using latent semantic an and summary evaluation. In In Proceedings ISIM '04. 2004. S. 93-100.
        #    LexRank - Unsupervised approach inspired by algorithms PageRank and HITS, reference
        #    TextRank - some sort of combination of a few resources that I found on the internet. I really don't remember the sources. Probably Wikipedia and some papers in 1st page of Google :)"""
        LANGUAGE = "english"
        #parser = HtmlParser.from_url(url, Tokenizer(LANGUAGE))
        if parser_option == 'file':
            parser = PlaintextParser.from_file(input, Tokenizer(LANGUAGE))
        elif parser_option == 'string':
            parser = PlaintextParser.from_string(input, Tokenizer(LANGUAGE))
        stemmer = Stemmer(LANGUAGE)
        
        summarizer = Summarizer(stemmer)
        summarizer.stop_words = get_stop_words(LANGUAGE)
        summary = []
        for sentence in summarizer(parser.document, SENTENCES_COUNT):
            summary.append(sentence)
        return summary
          
    def summarize_save(text, length, destination):   
        """ summarized text string and saves in destination file"""
        possible_title = text.split('\n')[0]
        possible_title = ''.join(c for c in possible_title if valid_xml_char_ordinal(c))
    
        summary = use_sumy(text, length, 'textrank', 'string')
#        with safe_open_a(destination) as f:
#            f.writelines('POSSIBLE TITLE: '+possible_title+'\n\n')
#            for line in summary:
#                f.writelines(str(line)+'\n\n')
#            f.writelines('\nFile Path: ' + path+'\n\n\n')
#            f.close()
        document.add_heading('Possible Title: '+possible_title, level=1)   
        for line in summary:
            line = str(line)
            line = unicode(line, errors='ignore')
            clean = ''.join(c for c in line if valid_xml_char_ordinal(c))
            document.add_paragraph(clean)
        path_string = path.replace("\\","/")
        document.add_paragraph('File Path: ' + path_string)
        if destination.endswith('.docx'):
            document.save(destination)
        else:
            document.save(destination+'.docx')

    fileTypes = ('.pdf', '.txt', '.docx', '.htm', '.html', 'htm', '.pptx')
    document = Document()
    if sum_type == 'directory':
        paths, files = recursive_glob(startpath, fileTypes)
        for idx, path in enumerate(paths):
            path = paths[idx]
            if files[idx].endswith('.pdf'):
                full_text = pdf2txt.convert_pdf_to_txt(path)
            elif files[idx].endswith('docx'):
                full_text = docx2txt.get_docx_text(path)
            else:
                full_text = textract.process(path)
            summarize_save(full_text, length, destination) 
    else:
        for path in fileList:
            if path.endswith('.pdf'):
                full_text = pdf2txt.convert_pdf_to_txt(path)
            else:
                full_text = textract.process(path)
            summarize_save(full_text, length, destination) 